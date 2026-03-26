import os
from dotenv import load_dotenv
load_dotenv()

import streamlit as st
import json
import io
from datetime import datetime
import docx as python_docx
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_CENTER

st.set_page_config(page_title="Manuscript Editor", page_icon="📖", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,400;0,600;1,400;1,600&family=DM+Sans:wght@300;400;500&display=swap');
html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
.block-container { padding-top: 2rem !important; max-width: 1100px !important; }
.main-title { font-family: 'Cormorant Garamond', serif; font-size: 3rem; font-weight: 400; font-style: italic; text-align: center; letter-spacing: 0.02em; margin-bottom: 0.3rem; }
.main-sub { text-align: center; color: #999; font-size: 0.85rem; font-weight: 300; margin-bottom: 2.5rem; letter-spacing: 0.05em; text-transform: uppercase; }
.edit-card { background: #fff; border: 1px solid #f0ece6; border-radius: 16px; padding: 1.6rem 2rem; margin-bottom: 1.2rem; border-left: 4px solid #c9a96e; }
.edit-card-header { display: flex; align-items: center; justify-content: space-between; margin-bottom: 1rem; }
.edit-card-title { font-family: 'Cormorant Garamond', serif; font-size: 1.25rem; font-style: italic; color: #2a2118; }
.score-pill { font-size: 0.75rem; font-weight: 500; padding: 4px 14px; border-radius: 20px; background: #fdf6ec; color: #a07840; border: 1px solid #e8d5b0; }
.headline-text { font-size: 0.9rem; color: #7a6a55; font-style: italic; margin-bottom: 1rem; padding-bottom: 1rem; border-bottom: 1px solid #f5f0e8; }
.feedback-text { font-size: 0.92rem; line-height: 1.9; color: #3a3028; white-space: pre-wrap; }
.revised-block { background: #fdf8f2; border: 1px solid #e8d5b0; border-radius: 10px; padding: 1rem 1.25rem; margin-top: 1rem; }
.revised-label { font-size: 0.72rem; text-transform: uppercase; letter-spacing: 0.1em; color: #a07840; margin-bottom: 0.5rem; font-weight: 500; }
.provider-bar { display: flex; align-items: center; justify-content: center; gap: 12px; margin-bottom: 2rem; flex-wrap: wrap; }
.provider-chip { font-size: 0.8rem; padding: 6px 16px; border-radius: 20px; border: 1px solid #e8e0d5; color: #7a6a55; cursor: pointer; background: white; }
.provider-chip.active { background: #2a2118; color: #f5e6c8; border-color: #2a2118; }
.stButton > button { background: #2a2118 !important; color: #f5e6c8 !important; border: none !important; border-radius: 10px !important; font-family: 'DM Sans', sans-serif !important; font-weight: 400 !important; font-size: 0.95rem !important; padding: 0.75rem 2rem !important; letter-spacing: 0.03em !important; width: 100% !important; }
.stButton > button:hover { background: #c9a96e !important; color: #2a2118 !important; }
div[data-testid="stTextInput"] input { border-radius: 8px !important; border: 1px solid #e8e0d5 !important; font-family: 'DM Sans', sans-serif !important; }
div[data-testid="stTextArea"] textarea { border-radius: 8px !important; border: 1px solid #e8e0d5 !important; font-family: 'DM Sans', sans-serif !important; font-size: 0.92rem !important; line-height: 1.8 !important; }
.summary-bar { display: grid; grid-template-columns: repeat(4,1fr); gap: 12px; margin-bottom: 2rem; }
.summary-item { text-align: center; padding: 1rem; background: #fdf8f2; border-radius: 12px; border: 1px solid #f0e8d8; }
.summary-score { font-family: 'Cormorant Garamond', serif; font-size: 2rem; font-style: italic; color: #2a2118; }
.summary-label { font-size: 0.72rem; text-transform: uppercase; letter-spacing: 0.08em; color: #a07840; margin-top: 2px; }
.chapter-item { padding: 0.9rem 1.2rem; border: 1px solid #f0e8d8; border-radius: 10px; margin-bottom: 8px; background: #fdfaf6; display: flex; align-items: center; justify-content: space-between; }
.key-input-wrap { background: #fdf8f2; border: 1px solid #e8d5b0; border-radius: 12px; padding: 1.2rem 1.5rem; margin-bottom: 1.5rem; }
.key-label { font-size: 0.78rem; text-transform: uppercase; letter-spacing: 0.08em; color: #a07840; margin-bottom: 0.5rem; font-weight: 500; }
</style>
""", unsafe_allow_html=True)

for k, v in [("history", []), ("api_keys", {}), ("provider", "Claude")]:
    if k not in st.session_state:
        st.session_state[k] = v

# Auto-load keys from .env file so you never have to paste them again
if not st.session_state.api_keys.get("Claude"):
    v = os.getenv("ANTHROPIC_API_KEY","")
    if v: st.session_state.api_keys["Claude"] = v
if not st.session_state.api_keys.get("ChatGPT"):
    v = os.getenv("OPENAI_API_KEY","")
    if v: st.session_state.api_keys["ChatGPT"] = v
if not st.session_state.api_keys.get("Gemini"):
    v = os.getenv("GEMINI_API_KEY","")
    if v: st.session_state.api_keys["Gemini"] = v

PROVIDERS = {
    "Claude":   {"lib": "anthropic", "models": ["claude-opus-4-5","claude-sonnet-4-5"], "placeholder": "sk-ant-...", "url": "console.anthropic.com"},
    "ChatGPT":  {"lib": "openai",    "models": ["gpt-4o","gpt-4o-mini"],               "placeholder": "sk-...",     "url": "platform.openai.com/api-keys"},
    "Gemini":   {"lib": "gemini",    "models": ["gemini-2.5-flash","gemini-2.0-flash-lite","gemini-1.5-flash"],   "placeholder": "AIza...",    "url": "aistudio.google.com"},
}

def call_ai(provider, key, model, prompt):
    if provider == "Claude":
        import anthropic
        c = anthropic.Anthropic(api_key=key)
        return c.messages.create(model=model, max_tokens=3000, messages=[{"role":"user","content":prompt}]).content[0].text.strip()
    elif provider == "ChatGPT":
        from openai import OpenAI
        c = OpenAI(api_key=key)
        return c.chat.completions.create(model=model, max_tokens=3000, messages=[{"role":"user","content":prompt}]).choices[0].message.content.strip()
    elif provider == "Gemini":
        import google.generativeai as genai
        genai.configure(api_key=key)
        m = genai.GenerativeModel(model)
        r = m.generate_content(prompt)
        return r.text.strip()

def parse_json(raw):
    raw = raw.replace("```json","").replace("```","").strip()
    return json.loads(raw[raw.index("{"):raw.rindex("}")+1])

# ── HEADER ──────────────────────────────────────────────────────────────────
st.markdown('<h1 class="main-title">The Manuscript Editor</h1>', unsafe_allow_html=True)
st.markdown('<p class="main-sub">Professional editorial feedback for publication-ready fiction</p>', unsafe_allow_html=True)

tab_edit, tab_history, tab_compare = st.tabs(["✍️  Edit", "📋  History", "↔️  Compare"])

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — EDITOR
# ══════════════════════════════════════════════════════════════════════════════
with tab_edit:

    # ── API Key + Provider row ───────────────────────────────────────────────
    st.markdown('<div class="key-input-wrap">', unsafe_allow_html=True)
    kc1, kc2, kc3 = st.columns([1.2, 2, 1])
    with kc1:
        st.markdown('<div class="key-label">AI Provider</div>', unsafe_allow_html=True)
        provider = st.selectbox("AI provider", list(PROVIDERS.keys()), label_visibility="collapsed",
            index=list(PROVIDERS.keys()).index(st.session_state.provider))
        st.session_state.provider = provider
        model = st.selectbox("Model", PROVIDERS[provider]["models"], label_visibility="collapsed")
    with kc2:
        st.markdown('<div class="key-label">API Key</div>', unsafe_allow_html=True)
        saved_key = st.session_state.api_keys.get(provider, "")
        key_input = st.text_input("API key", value=saved_key, type="password",
            placeholder=PROVIDERS[provider]["placeholder"], label_visibility="collapsed")
        if key_input:
            st.session_state.api_keys[provider] = key_input
    with kc3:
        st.markdown('<div class="key-label">Get your key</div>', unsafe_allow_html=True)
        url = PROVIDERS[provider]["url"]
        st.markdown(f'<a href="https://{url}" target="_blank" style="font-size:0.82rem;color:#a07840;">{url}</a>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Input ────────────────────────────────────────────────────────────────
    ic1, ic2 = st.columns([1,1], gap="large")
    with ic1:
        chapter_label = st.text_input("Chapter label", placeholder="e.g. Chapter 12 — The Confrontation")
        method = st.radio("Input method", ["Paste text", "Upload file"], horizontal=True, label_visibility="collapsed")

        text = ""
        if method == "Paste text":
            text = st.text_area("Chapter text", height=420, placeholder="Paste your chapter here…", label_visibility="collapsed")
        else:
            f = st.file_uploader("Upload file", type=["txt","docx"], label_visibility="collapsed")
            if f:
                text = f.read().decode("utf-8","ignore") if f.name.endswith(".txt") else \
                    "\n".join(p.text for p in python_docx.Document(io.BytesIO(f.read())).paragraphs if p.text.strip())
                st.success(f"Loaded {f.name} — {len(text.split()):,} words")
                st.text_area("Preview", text[:600], height=120, disabled=True)

        if text.strip():
            st.caption(f"{len(text.split()):,} words · {len(text):,} characters")

        go = st.button("Edit this chapter →")

    # ── Output ───────────────────────────────────────────────────────────────
    with ic2:
        if go:
            active_key = st.session_state.api_keys.get(provider, "")
            if not active_key:
                st.error(f"Paste your {provider} API key above first.")
            elif not text.strip():
                st.error("Paste or upload a chapter on the left.")
            else:
                prompt = f"""You are a senior literary editor at a major publishing house, specialising in commercial fiction — Romance, Drama, Women's Fiction.

A novelist has submitted the following chapter for editorial review ahead of publication. Read it carefully, understand its genre, tone, POV, and characters from the text itself — do not ask for any of this information.

Your job is to give the kind of deep, honest, constructive editorial feedback that will make this chapter publication-ready. Be specific and reference actual lines from the text. Be encouraging but ruthlessly honest — this author is serious about getting published.

Return ONLY a raw JSON object in exactly this format:
{{
  "overall": {{
    "score": 8,
    "verdict": "One punchy sentence summarising the chapter's biggest strength and biggest opportunity."
  }},
  "line_edits": {{
    "score": 7,
    "headline": "One sentence on the prose quality overall.",
    "notes": "Detailed line-level feedback. Reference specific sentences. Flag passive voice, awkward phrasing, repetitive words, over-writing, under-writing. Show 2-3 actual lines with suggested rewrites.",
    "rewrite_example": "Pick the weakest sentence and show a better version."
  }},
  "structure": {{
    "score": 8,
    "headline": "One sentence on how well the scene is constructed.",
    "notes": "Does the scene have a clear purpose? Does it open strongly and end with momentum? Where does it drag? What should be cut or expanded? Be specific about which paragraphs.",
    "rewrite_example": "Suggest how the opening or closing of the scene could be stronger."
  }},
  "character": {{
    "score": 9,
    "headline": "One sentence on character authenticity.",
    "notes": "Are the characters behaving consistently and authentically? Is the dialogue natural and distinctive per character? Is the emotional interiority earning its space on the page? Flag any moments that feel false.",
    "rewrite_example": "Show a dialogue or interiority line rewritten to feel more authentic."
  }},
  "emotional_impact": {{
    "score": 8,
    "headline": "One sentence on how emotionally effective this chapter is.",
    "notes": "This is Romance and Drama — emotional resonance is everything. Is the tension landing? Are the stakes clear? Does the reader feel what the character feels? What moments are working and what are falling flat?",
    "rewrite_example": "Rewrite a key emotional moment to hit harder."
  }}
}}

Chapter{f' ({chapter_label})' if chapter_label else ''}:
---
{text[:7000]}
---"""

                with st.spinner("Reading your chapter…"):
                    try:
                        raw    = call_ai(provider, active_key, model, prompt)
                        result = parse_json(raw)

                        overall = result.get("overall", {})
                        sections = {
                            "line_edits":       ("Line Edits & Prose",    "#c9a96e"),
                            "structure":        ("Scene Structure",        "#8fa882"),
                            "character":        ("Character & Dialogue",   "#8ba5b8"),
                            "emotional_impact": ("Emotional Impact",       "#b88fa0"),
                        }

                        # Summary bar
                        scores = {k: result[k]["score"] for k in sections if k in result}
                        st.markdown('<div class="summary-bar">', unsafe_allow_html=True)
                        for k, (label, color) in sections.items():
                            s = scores.get(k, "—")
                            st.markdown(f'''<div class="summary-item">
                                <div class="summary-score">{s}</div>
                                <div class="summary-label">{label}</div>
                            </div>''', unsafe_allow_html=True)
                        st.markdown('</div>', unsafe_allow_html=True)

                        if overall:
                            st.markdown(f'<div style="text-align:center;font-family:Cormorant Garamond,serif;font-size:1.1rem;font-style:italic;color:#7a6a55;margin-bottom:1.5rem;padding:1rem;background:#fdf8f2;border-radius:10px;border:1px solid #e8d5b0;">{overall.get("verdict","")}</div>', unsafe_allow_html=True)

                        for k, (label, color) in sections.items():
                            if k not in result: continue
                            e = result[k]
                            st.markdown(f'''<div class="edit-card" style="border-left-color:{color};">
  <div class="edit-card-header">
    <span class="edit-card-title">{label}</span>
    <span class="score-pill">{e.get("score","—")} / 10</span>
  </div>
  <div class="headline-text">{e.get("headline","")}</div>
  <div class="feedback-text">{e.get("notes","")}</div>
  <div class="revised-block">
    <div class="revised-label">Suggested rewrite</div>
    <div class="feedback-text">{e.get("rewrite_example","")}</div>
  </div>
</div>''', unsafe_allow_html=True)

                        st.session_state.history.append({
                            "label":     chapter_label or f"Chapter {len(st.session_state.history)+1}",
                            "timestamp": datetime.now().strftime("%d %b %Y, %H:%M"),
                            "provider":  provider, "model": model,
                            "text":      text, "result": result, "scores": scores,
                        })

                        # Export
                        st.markdown("---")
                        ex1, ex2 = st.columns(2)
                        with ex1:
                            wd = python_docx.Document()
                            wd.add_heading("Editorial Feedback", 0)
                            wd.add_paragraph(f"{chapter_label or 'Chapter'} · {datetime.now().strftime('%d %b %Y')} · {provider}/{model}")
                            if overall:
                                wd.add_paragraph(overall.get("verdict",""))
                            for k, (label, _) in sections.items():
                                if k not in result: continue
                                e = result[k]
                                wd.add_heading(f"{label} — {e.get('score')}/10", level=2)
                                wd.add_paragraph(e.get("headline",""), style="Intense Quote")
                                wd.add_paragraph(e.get("notes",""))
                                wd.add_heading("Suggested rewrite", level=3)
                                wd.add_paragraph(e.get("rewrite_example",""))
                            buf = io.BytesIO(); wd.save(buf); buf.seek(0)
                            st.download_button("⬇ Download as Word", buf,
                                file_name=f"edit_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        with ex2:
                            pb = io.BytesIO()
                            dp = SimpleDocTemplate(pb, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
                            ts = ParagraphStyle("t", fontName="Helvetica-Bold", fontSize=18, spaceAfter=6, alignment=TA_CENTER)
                            ss = ParagraphStyle("s", fontName="Helvetica", fontSize=9, spaceAfter=14, alignment=TA_CENTER, textColor=colors.grey)
                            h2 = ParagraphStyle("h", fontName="Helvetica-Bold", fontSize=13, spaceBefore=16, spaceAfter=4)
                            h3 = ParagraphStyle("h3", fontName="Helvetica-Bold", fontSize=10, spaceBefore=10, spaceAfter=3, textColor=colors.HexColor("#a07840"))
                            bs = ParagraphStyle("b", fontName="Helvetica", fontSize=10, leading=16, spaceAfter=10)
                            story = [Paragraph("Editorial Feedback", ts),
                                     Paragraph(f"{chapter_label or 'Chapter'} · {datetime.now().strftime('%d %b %Y')}", ss),
                                     HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e8d5b0"), spaceAfter=14)]
                            if overall:
                                story.append(Paragraph(f"<i>{overall.get('verdict','')}</i>", bs))
                                story.append(Spacer(1,8))
                            for k, (label, _) in sections.items():
                                if k not in result: continue
                                e = result[k]
                                story += [Paragraph(f"{label} — {e.get('score')}/10", h2),
                                          Paragraph(f"<i>{e.get('headline','')}</i>", bs),
                                          Paragraph(e.get("notes","").replace("\n","<br/>"), bs),
                                          Paragraph("Suggested rewrite", h3),
                                          Paragraph(e.get("rewrite_example","").replace("\n","<br/>"), bs),
                                          Spacer(1,6)]
                            dp.build(story); pb.seek(0)
                            st.download_button("⬇ Download as PDF", pb,
                                file_name=f"edit_{datetime.now().strftime('%Y%m%d')}.pdf", mime="application/pdf")

                    except json.JSONDecodeError:
                        st.error("Unexpected response format — please try again.")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
        else:
            st.markdown("""<div style="height:420px;display:flex;align-items:center;justify-content:center;color:#c8b89a;font-family:'Cormorant Garamond',serif;font-size:1.3rem;font-style:italic;border:1px dashed #e8d5b0;border-radius:12px;">
Your editorial feedback will appear here</div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — HISTORY
# ══════════════════════════════════════════════════════════════════════════════
with tab_history:
    if not st.session_state.history:
        st.info("No chapters edited yet.")
    else:
        total = st.session_state.history.__len__()
        st.caption(f"{total} chapter{'s' if total!=1 else ''} edited")
        for h in reversed(st.session_state.history):
            with st.expander(f"📄 {h['label']}  ·  {h['timestamp']}  ·  {h.get('provider','')}"):
                sc = h.get("scores", {})
                c1,c2,c3,c4 = st.columns(4)
                for col, k in zip([c1,c2,c3,c4], ["line_edits","structure","character","emotional_impact"]):
                    labels = {"line_edits":"Prose","structure":"Structure","character":"Character","emotional_impact":"Emotion"}
                    col.metric(labels[k], f"{sc.get(k,'—')}/10" if sc.get(k) else "—")
                st.markdown("---")
                for k, (label, _) in [("line_edits",("Line Edits","")),(("structure"),("Structure","")),(("character"),("Character","")),(("emotional_impact"),("Emotional Impact",""))]:
                    if k in h["result"]:
                        st.markdown(f"**{label}** — _{h['result'][k].get('headline','')}_")
                        st.write(h["result"][k].get("notes",""))

        if st.button("⬇ Download full history as Word"):
            wd = python_docx.Document()
            wd.add_heading("Full Editorial Report", 0)
            wd.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y')}")
            for h in st.session_state.history:
                wd.add_page_break()
                wd.add_heading(h["label"], level=1)
                for k, entry in h["result"].items():
                    if k == "overall": continue
                    labels = {"line_edits":"Line Edits","structure":"Structure","character":"Character","emotional_impact":"Emotional Impact"}
                    wd.add_heading(f"{labels.get(k,k)} — {entry.get('score')}/10", level=2)
                    wd.add_paragraph(entry.get("notes",""))
            buf = io.BytesIO(); wd.save(buf); buf.seek(0)
            st.download_button("Click to download", buf, file_name="full_editorial_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — COMPARE
# ══════════════════════════════════════════════════════════════════════════════
with tab_compare:
    if not st.session_state.history:
        st.info("Edit a chapter first to use the compare view.")
    else:
        sel = st.selectbox("Select chapter", [h["label"] for h in st.session_state.history])
        h   = next(x for x in st.session_state.history if x["label"] == sel)
        l, r = st.columns(2, gap="large")
        with l:
            st.markdown("**Your original text**")
            st.text_area("", h["text"][:5000], height=600, disabled=True, label_visibility="collapsed")
        with r:
            st.markdown("**Editorial notes**")
            out = ""
            labels = {"line_edits":"Line Edits","structure":"Structure","character":"Character & Dialogue","emotional_impact":"Emotional Impact"}
            for k, label in labels.items():
                if k in h["result"]:
                    e = h["result"][k]
                    out += f"── {label} ({e.get('score')}/10) ──\n{e.get('headline','')}\n\n{e.get('notes','')}\n\nSuggested rewrite:\n{e.get('rewrite_example','')}\n\n"
            st.text_area("", out, height=600, disabled=True, label_visibility="collapsed")
